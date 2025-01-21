import docx
import os
import re
import cv2
import numpy as np

class DocxManager:
    def __init__(self):
        pass

    def replace_words(self, path, old, new):
        """
        :param path:文件路径
        :param old:需要替换的keyword
        :param new:新的替换后的keyword
        """
        if path.endswith(".docx"):
            # 不支持读取doc格式的文件
            doc = docx.Document(path)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = run.text.replace(old, new)
                doc.save(path)
        else:
            raise ValueError("只支持docx文件格式!")
        
    
    def change_forward(self, word_path, result_path):
        """
        更改Word方向
        :param word_path: word文件路径
        :param result_path: 结果文件路径

        use:
        path = 'Robot'
        spam=os.listdir(path)
        os.chdir(path)
        for i in spam:
            if i.endswith('.docx'):
                get_pictures(str(i),os.getcwd())
        """
        if not os.path.exists(result_path):
            os.makedirs(result_path)
        doc = docx.Document(word_path)
        for section in doc.sections:
            # 交替宽高
            section.page_width,section.page_height = section.page_height ,section.page_width
        # 保存为新文件
        doc.save(os.path.join(result_path,word_path)) 

    def get_pictures(self, word_path, result_path):
        """
        图片提取
        :param word_path: word路径
        :result_path: 保存路径
        :return: 
        """
        # 创建保存路径
        if not os.path.exists(result_path):
            os.makedirs(result_path)
        # 读取文件
        doc = docx.Document(word_path)

        # 获取图片
        dict_rel = doc.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:            
                img_name = re.findall("/(.*)", rel.target_ref)[0]
                word_name = os.path.splitext(word_path)[0]
                if os.sep in word_name:
                    new_name = word_name.split('\\')[-1]
                else:
                    new_name = word_name.split('/')[-1]
                # cv2获取图片大小
                imgdata = np.frombuffer(rel.target_part.blob,np.uint8)
                img_cv = cv2.imdecode(imgdata,cv2.IMREAD_ANYCOLOR)
                img_name = '{}-{}-{}-{}'.format(new_name,img_cv.shape[0],img_cv.shape[1],img_name)
                # 直接二进制写入兼容性比使用CV2的保存图片好
                with open(f'{result_path}/{img_name}','wb') as f:
                    f.write(rel.target_part.blob)
            else:
                pass