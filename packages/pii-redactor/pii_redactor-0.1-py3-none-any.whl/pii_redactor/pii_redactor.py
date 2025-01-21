# pii_redactor.py
import re
import json
import logging
import spacy
from pathlib import Path
from typing import Dict, List, Union, Optional, Any
from collections import defaultdict
from dataclasses import dataclass, asdict
import pdfplumber
from docx import Document
from phonenumbers import PhoneNumberMatcher
from luhn import verify
import yaml

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class RedactionRecord:
    start: int
    end: int
    original: str
    pii_type: str

class PIIConfig:
    def __init__(self, config_path: str = None):
        self.default_config = {
            'patterns': {
                'email': r'\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b',
                'ssn': r'\b(?!000|666|9\d{2})\d{3}-(?!00)\d{2}-(?!0000)\d{4}\b',
                'credit_card': r'\b(?:\d[ -]*?){13,19}\b',
                'phone': r'''
                    (?:\+?(\d{1,3}))?                # Country code
                    (?:[-.\s]?\(?\d{1,4}\)?)?        # Area code
                    (?:[-.\s]+\d{2,5}){2,}           # Local number
                    \b
                ''',
                'org_keywords': r'\b(Company|Corp|LLC|Inc|Ltd)\b'
            },
            'locales': {
                'en_US': {
                    'ner_model': 'en_core_web_sm',
                    'honorifics': ['Dr.', 'Mr.', 'Mrs.', 'Ms.', 'Prof.'],
                    'patterns': {}
                },
                'es_ES': {
                    'ner_model': 'es_core_news_sm',
                    'honorifics': ['Dr.', 'Dra.', 'Sr.', 'Sra.', 'Don', 'Doña'],
                    'patterns': {
                        'dni': r'(DNI|NIE)\s*:?\s*[\dX-Z]{8,9}',
                        'iban': r'ES\d{2}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{4}\s?\d{4}',
                        'org_keywords': r'\b(Empresa|Compañía|Corporación)\b'
                    }
                }
            },
            'supported_formats': ['.txt', '.docx', '.pdf']
        }
        
        self.config = self.default_config
        if config_path:
            self.load_config(config_path)

    def load_config(self, path: str):
        try:
            with open(path, 'r') as f:
                user_config = yaml.safe_load(f) or {}
                self._deep_update(self.config, user_config)
        except Exception as e:
            logger.error(f"Failed to load config: {str(e)}")

    def _deep_update(self, base: Dict[Any, Any], update: Dict[Any, Any]):
        for k, v in update.items():
            if isinstance(v, dict) and k in base:
                self._deep_update(base[k], v)
            else:
                base[k] = v

class FileHandler:
    @staticmethod
    def read(file_path: Path) -> str:
        if file_path.suffix == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                return '\n'.join([page.extract_text() for page in pdf.pages])
        elif file_path.suffix == '.docx':
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        else:
            return file_path.read_text(encoding='utf-8')

    @staticmethod
    def write(content: str, output_path: Path):
        if output_path.suffix == '.pdf':
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            c = canvas.Canvas(str(output_path), pagesize=letter)
            text = c.beginText(40, 750)
            text.textLines(content)
            c.drawText(text)
            c.save()
        elif output_path.suffix == '.docx':
            doc = Document()
            doc.add_paragraph(content)
            doc.save(output_path)
        else:
            output_path.write_text(content, encoding='utf-8')

class PIIRedactor:
    def __init__(self, locale: str = 'en_US', config: PIIConfig = None):
        self.config = config or PIIConfig()
        self.locale = locale
        self.redaction_char = '█'
        self.redaction_log: List[RedactionRecord] = []
        self.organization_suffixes = {'LLC', 'Inc', 'Ltd', 'Corp', 'Co', 'AG', 'SL', 'SA'}
        
        # Load NLP model
        try:
            model_name = self.config.config['locales'][locale]['ner_model']
            self.nlp = spacy.load(model_name)
        except KeyError:
            logger.error(f"No model configured for locale {locale}")
            raise
        except OSError:
            logger.error(f"Model {model_name} not installed. Run: python -m spacy download {model_name}")
            raise

    def detect_pii(self, text: str) -> Dict[str, List[str]]:
        results = defaultdict(list)
        logger.debug(f"Processing text: {text[:200]}...")
        locale_config = self.config.config['locales'].get(self.locale, {})
        
        # Regex patterns detection
        patterns = {**self.config.config['patterns'], **locale_config.get('patterns', {})}
        for pii_type, pattern in patterns.items():
            try:
                compiled = re.compile(pattern, re.VERBOSE | re.IGNORECASE)
                for match in compiled.finditer(text):
                    results[pii_type].append(match.group())
            except Exception as e:
                logger.error(f"Pattern error for {pii_type}: {str(e)}")

        # NER detection
        doc = self.nlp(text)
        for ent in doc.ents:
            if ent.label_ in ('PER', 'PERSON'):
                self._process_person_entity(ent, results, doc)
            elif ent.label_ == 'ORG':
                self._process_org_entity(ent, results, doc)
        
        logger.debug(f"Detected PII: {dict(results)}")
        return dict(results)

    def _process_person_entity(self, ent, results, doc):
        """Handle names with titles and honorifics"""
        locale_config = self.config.config['locales'].get(self.locale, {})
        honorifics = locale_config.get('honorifics', [])
        
        # Capture honorifics within 3 tokens before the entity
        for i in range(max(0, ent.start-3), ent.start):
            if doc[i].text in honorifics:
                full_span = doc[i:ent.end]
                full_name = full_span.text
                results['names'].append(full_name)
                return
        
        # Capture multi-word names without honorifics
        if len(ent.text.split()) >= 2 and any(t.text.istitle() for t in ent):
            results['names'].append(ent.text)

    def _process_org_entity(self, ent, results, doc):
        """Handle organizations with suffixes and keywords"""
        org_tokens = [ent.text]
        current_idx = ent.end
        
        # Capture adjacent organization suffixes and acronyms
        while current_idx < len(doc):
            next_token = doc[current_idx]
            if next_token.text in self.organization_suffixes or next_token.text.isupper():
                org_tokens.append(next_token.text)
                current_idx += 1
            else:
                break
        
        # Capture preceding organization keywords
        for i in range(max(0, ent.start-2), ent.start):
            if doc[i].text.lower() in {'at', 'de', 'of', 'en'}:
                org_tokens.insert(0, doc[i].text)
        
        results['organizations'].append(" ".join(org_tokens))

    def redact_text(self, text: str, pii_data: Dict[str, List[str]]) -> str:
        text_list = list(text)
        redact_ranges = []
        
        for pii_type, matches in pii_data.items():
            for match in matches:
                start_idx = 0
                while True:
                    start = text.find(match, start_idx)
                    if start == -1:
                        break
                    end = start + len(match)
                    redact_ranges.append((start, end, pii_type))
                    self.redaction_log.append(RedactionRecord(
                        start=start,
                        end=end,
                        original=text[start:end],
                        pii_type=pii_type
                    ))
                    start_idx = end
        
        # Merge overlapping ranges
        redact_ranges.sort()
        merged = []
        for start, end, p_type in redact_ranges:
            if merged and start <= merged[-1][1]:
                merged[-1] = (merged[-1][0], max(end, merged[-1][1]), p_type)
            else:
                merged.append((start, end, p_type))
        
        # Apply redaction
        for start, end, _ in merged:
            for i in range(start, end):
                if i < len(text_list) and text_list[i] not in ('\n', '\r'):
                    text_list[i] = self.redaction_char
        
        return ''.join(text_list)

    def reverse_redaction(self, redacted_text: str, log: List[RedactionRecord]) -> str:
        text_list = list(redacted_text)
        for record in sorted(log, key=lambda x: -x.start):
            original = list(record.original)
            for i in range(record.start, record.end):
                if i < len(text_list):
                    try:
                        text_list[i] = original[i - record.start]
                    except IndexError:
                        continue
        return ''.join(text_list)

    def process_file(self, input_path: Union[str, Path], output_path: Union[str, Path]) -> Optional[Dict]:
        try:
            input_path = Path(input_path)
            output_path = Path(output_path)

            if input_path.suffix not in self.config.config['supported_formats']:
                raise ValueError(f"Unsupported file format: {input_path.suffix}")

            # Read file
            text = FileHandler.read(input_path)
            
            # Process content
            pii_data = self.detect_pii(text)
            redacted = self.redact_text(text, pii_data)
            
            # Write output
            FileHandler.write(redacted, output_path)
            
            # Save redaction log
            log_path = output_path.with_suffix('.log.json')
            log_data = [asdict(record) for record in self.redaction_log]
            log_path.write_text(json.dumps(log_data))
            
            return pii_data

        except Exception as e:
            logger.error(f"Processing failed: {str(e)}", exc_info=True)
            return None

def main():
    """Command-line interface"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Advanced PII Redaction Tool')
    parser.add_argument('input', help='Input file path')
    parser.add_argument('output', help='Output file path')
    parser.add_argument('--locale', default='en_US', 
                       help='Locale for localization (e.g., en_US, es_ES)')
    parser.add_argument('--config', help='Path to custom config YAML file')
    parser.add_argument('--reverse', metavar='LOG_FILE', 
                       help='Reverse redaction using log file')
    parser.add_argument('--verbose', '-v', action='store_true', 
                       help='Enable verbose output')

    args = parser.parse_args()

    if args.verbose:
        logger.setLevel(logging.DEBUG)

    try:
        if args.reverse:
            # Reverse mode
            redactor = PIIRedactor()
            log_data = json.loads(Path(args.reverse).read_text())
            log = [RedactionRecord(**r) for r in log_data]
            
            content = FileHandler.read(Path(args.input))
            restored = redactor.reverse_redaction(content, log)
            FileHandler.write(restored, Path(args.output))
            
            print(f"Successfully reversed redaction to {args.output}")
        else:
            # Redaction mode
            config = PIIConfig(args.config) if args.config else None
            redactor = PIIRedactor(locale=args.locale, config=config)
            result = redactor.process_file(args.input, args.output)
            
            if result:
                print(f"Redacted file saved to {args.output}")
                print("Redaction log:", Path(args.output).with_suffix('.log.json'))
                print("\nDetected PII:")
                for pii_type, items in result.items():
                    print(f"  {pii_type}: {len(items)}")
            else:
                print("Redaction failed")
                exit(1)

    except Exception as e:
        logger.error(f"Fatal error: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()