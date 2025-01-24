from typing import List
from bcm.models import LayoutModel
from bcm.layout_manager import process_layout
from bcm.settings import Settings

from docx import Document

def _format_description_word(description: str, document) -> None:
    """Format description text for Word, handling newlines."""
    if not description:
        return  # Don't add a paragraph for no description

    # Add each non-empty line as a separate paragraph
    for line in description.split('\n'):
        if line.strip():  # Check if the line has content after removing whitespace
            document.add_paragraph(line)

def _process_node_word(document: Document, node: LayoutModel, level: int = 0) -> None:
    """Process a node and its children recursively to generate Word content."""

    # Add header with proper heading level
    if level == 0:
        if node.name:
            document.add_paragraph(node.name, style='Title')
    elif level == 1:
        if node.name:
            document.add_paragraph(node.name, style='Heading 1')
    elif level == 2:
        if node.name:
            document.add_paragraph(node.name, style='Heading 2')
    elif level == 3:
        if node.name:
            document.add_paragraph(node.name, style='Heading 3')
    elif level == 4:
        if node.name:
            document.add_paragraph(node.name, style='Heading 4')
    else:
        if node.name:
            document.add_paragraph(node.name, style='Heading 5') # Or use a bullet point if deeper

    # Add description if present
    if node.description:
        _format_description_word(node.description, document)

    # Process children if present
    if node.children:
        for child in node.children:
            _process_node_word(document, child, level + 1)

def export_to_word(model: LayoutModel, settings: Settings) -> Document:
    """Export the capability model to a Microsoft Word Document object.

    Args:
        model: The capability model to export.
        settings: Application settings.

    Returns:
        A docx.Document object representing the capability model in Word format.
    """
    document = Document()

    # Process layout first
    processed_model = process_layout(model, settings)

    # Process all nodes recursively
    _process_node_word(document, processed_model, level=0)

    return document