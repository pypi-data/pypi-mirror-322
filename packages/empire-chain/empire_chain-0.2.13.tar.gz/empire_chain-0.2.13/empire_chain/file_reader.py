from typing import Protocol
from pathlib import Path
import PyPDF2
import docx
import json
import csv
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
import os
import pickle

class FileReader(Protocol):
    def read(self, file_path: str) -> str:
        pass

class PDFReader(FileReader):
    def read(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

class DocxReader(FileReader):
    def read(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

class TxtReader(FileReader):
    def read(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

class JSONReader(FileReader):
    def read(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            return json.dumps(data, indent=2)

class CSVReader(FileReader):
    def read(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            for row in reader:
                text += ",".join(row) + "\n"
        return text

class GoogleDocsReader(FileReader):
    def __init__(self):
        """Initialize Google Drive reader with OAuth authentication."""
        self.SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
        self.creds = None
        self.token_path = 'token.pickle'
    
    def _get_credentials(self):
        """Gets valid user credentials from storage or initiates OAuth flow.
        
        Returns:
            Credentials, the obtained credential.
        """
        if os.path.exists(self.token_path):
            with open(self.token_path, 'rb') as token:
                self.creds = pickle.load(token)

        if not self.creds or not self.creds.valid:
            if self.creds and self.creds.expired and self.creds.refresh_token:
                self.creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(
                    'client_secrets.json',  # This will be created by the user
                    self.SCOPES,
                    redirect_uri='http://localhost:8080'
                )
                self.creds = flow.run_local_server(port=8080)
                
            with open(self.token_path, 'wb') as token:
                pickle.dump(self.creds, token)

        return self.creds

    def _extract_file_id(self, file_path: str) -> str:
        """Extract file ID from Google Drive URL or return as is if already an ID."""
        if 'drive.google.com' in file_path:
            if '/file/d/' in file_path:
                file_id = file_path.split('/file/d/')[1].split('/')[0]
            elif '/document/d/' in file_path:
                file_id = file_path.split('/document/d/')[1].split('/')[0]
            else:
                raise ValueError(
                    "Invalid Google Drive URL format. Please use the 'Share' link from Google Drive."
                )
            return file_id
        return file_path

    def read(self, file_path: str) -> str:
        """Reads a Google Drive file and returns its content as text.
        
        Args:
            file_path: The Google Drive file ID or URL
            
        Returns:
            str: The document content as text
            
        Raises:
            ValueError: If authentication fails or file cannot be accessed
        """
        try:
            creds = self._get_credentials()
        except Exception as e:
            raise ValueError(
                "Authentication failed. Please make sure you have client_secrets.json in your working directory. "
                "You can get it from Google Cloud Console > APIs & Services > Credentials > Create Credentials > OAuth Client ID. "
                f"Error: {str(e)}"
            )
            
        drive_service = build('drive', 'v3', credentials=creds)
        docs_service = build('docs', 'v1', credentials=creds)
        
        file_id = self._extract_file_id(file_path)
        
        try:
            file_metadata = drive_service.files().get(fileId=file_id, fields='mimeType').execute()
            mime_type = file_metadata['mimeType']
            
            if mime_type == 'application/vnd.google-apps.document':
                document = docs_service.documents().get(documentId=file_id).execute()
                text = ""
                for content in document.get('body').get('content'):
                    if 'paragraph' in content:
                        for element in content.get('paragraph').get('elements'):
                            if 'textRun' in element:
                                text += element.get('textRun').get('content')
                return text
                
            else:
                request = drive_service.files().get_media(fileId=file_id)
                file_content = io.BytesIO()
                downloader = MediaIoBaseDownload(file_content, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                
                file_content.seek(0)
                
                if mime_type == 'application/pdf':
                    reader = PyPDF2.PdfReader(file_content)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
                elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    doc = docx.Document(file_content)
                    return "\n".join(paragraph.text for paragraph in doc.paragraphs)
                elif mime_type == 'text/plain':
                    return file_content.read().decode('utf-8')
                else:
                    raise ValueError(f"Unsupported Google Drive file type: {mime_type}")
                    
        except Exception as e:
            raise ValueError(f"Could not read Google Drive file: {str(e)}")

class DocumentReader:
    def __init__(self):
        """Initialize document reader."""
        self.readers = {
            '.pdf': PDFReader(),
            '.docx': DocxReader(),
            '.txt': TxtReader(),
            '.json': JSONReader(),
            '.csv': CSVReader()
        }
        self.google_reader = GoogleDocsReader()
    
    def _is_google_drive_url(self, file_path: str) -> bool:
        """Check if the given path is a Google Drive URL."""
        return 'drive.google.com' in file_path
    
    def read(self, file_path: str) -> str:
        """Read content from various file types and return as text.
        
        Args:
            file_path: Path to the file to read or Google Drive URL
            
        Returns:
            str: Text content of the file
            
        Raises:
            ValueError: If file type is not supported or authentication fails
        """
        if self._is_google_drive_url(file_path):
            return self.google_reader.read(file_path)
            
        file_extension = Path(file_path).suffix.lower()
        if not file_extension:
            raise ValueError("File has no extension and is not a Google Drive URL")
            
        if file_extension not in self.readers:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return self.readers[file_extension].read(file_path)
    
    def supported_formats(self) -> list[str]:
        """Get list of supported file formats.
        
        Returns:
            list[str]: List of supported file extensions
        """
        return list(self.readers.keys()) 